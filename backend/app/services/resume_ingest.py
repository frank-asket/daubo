"""Extract plain text from resume uploads (PDF, DOCX, text, images via vision)."""

from __future__ import annotations

import base64
import io
import logging
from typing import TYPE_CHECKING

import httpx
from docx import Document
from PIL import Image
from pypdf import PdfReader

if TYPE_CHECKING:
    from app.config import Settings

logger = logging.getLogger("daubo")

MAX_UPLOAD_BYTES = 12 * 1024 * 1024
MAX_RESUME_CHARS = 500_000

_IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif"})


def _normalize_text(raw: str) -> str:
    t = raw.replace("\x00", " ").strip()
    if len(t) > MAX_RESUME_CHARS:
        t = t[:MAX_RESUME_CHARS]
    return t


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _normalize_text("\n".join(parts))


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return _normalize_text("\n".join(parts))


def _extract_plain(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return _normalize_text(data.decode(enc))
        except UnicodeDecodeError:
            continue
    return _normalize_text(data.decode("utf-8", errors="replace"))


def _pil_mime(fmt: str | None) -> str:
    if fmt == "PNG":
        return "image/png"
    if fmt in ("JPEG", "JPG"):
        return "image/jpeg"
    if fmt == "WEBP":
        return "image/webp"
    if fmt == "GIF":
        return "image/gif"
    return "image/png"


async def _vision_resume_text(settings: Settings, image_bytes: bytes, mime: str) -> str:
    if not settings.openrouter_api_key:
        raise ValueError(
            "Image resumes need OPENROUTER_API_KEY for vision transcription. "
            "Upload PDF/DOCX, or paste text instead."
        )
    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    payload: dict = {
        "model": settings.openrouter_vision_model,
        "temperature": settings.openrouter_temperature,
        "top_p": settings.openrouter_top_p,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Transcribe all readable text from this resume or CV image. "
                            "Output plain text only. Preserve approximate structure using "
                            "line breaks. No preamble or commentary."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                ],
            }
        ],
    }
    if settings.openrouter_top_k is not None:
        payload["top_k"] = settings.openrouter_top_k
    headers = {
        "Authorization": f"Bearer {settings.openrouter_api_key}",
        "HTTP-Referer": settings.openrouter_http_referer,
        "X-Title": settings.openrouter_app_title,
        "Content-Type": "application/json",
    }
    async with httpx.AsyncClient(timeout=120.0) as client:
        r = await client.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=payload,
        )
        if not r.is_success:
            logger.warning("OpenRouter vision error: %s %s", r.status_code, r.text[:500])
            r.raise_for_status()
        data = r.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError) as exc:
        raise ValueError("Vision model returned an unexpected response.") from exc
    text = content if isinstance(content, str) else str(content)
    return _normalize_text(text)


async def extract_resume_text(
    data: bytes,
    filename: str,
    content_type: str | None,
    settings: Settings,
) -> str:
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File too large (max {MAX_UPLOAD_BYTES // (1024 * 1024)} MB).")

    name = (filename or "upload").strip()
    lower = name.lower()
    ct = (content_type or "").split(";")[0].strip().lower()

    if lower.endswith(".pdf") or ct == "application/pdf":
        text = _extract_pdf(data)
    elif lower.endswith(".docx") or ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        text = _extract_docx(data)
    elif lower.endswith(".txt") or ct in ("text/plain",):
        text = _extract_plain(data)
    elif lower.endswith(".doc") or ct == "application/msword":
        raise ValueError("Legacy .doc is not supported. Save as .docx or PDF and upload again.")
    elif (
        any(lower.endswith(s) for s in _IMAGE_SUFFIXES)
        or ct.startswith("image/")
        or ct
        in (
            "image/png",
            "image/jpeg",
            "image/webp",
            "image/gif",
        )
    ):
        try:
            img = Image.open(io.BytesIO(data))
            img.load()
            mime = ct if ct.startswith("image/") else _pil_mime(img.format)
        except Exception as exc:
            raise ValueError("Could not read image file.") from exc
        text = await _vision_resume_text(settings, data, mime)
    else:
        if data[:4] == b"%PDF":
            text = _extract_pdf(data)
        elif data[:2] == b"PK":
            text = _extract_docx(data)
        elif ct.startswith("image/"):
            text = await _vision_resume_text(settings, data, ct)
        else:
            try:
                img = Image.open(io.BytesIO(data))
                img.load()
                mime = _pil_mime(img.format)
                text = await _vision_resume_text(settings, data, mime)
            except Exception as exc:
                raise ValueError(
                    "Unsupported file type. Use PDF, DOCX, TXT, or a common image (PNG, JPEG, WebP)."
                ) from exc

    if not text.strip():
        raise ValueError(
            "No text could be extracted. For scanned PDFs, try exporting as images or paste "
            "the text. For images, ensure OPENROUTER_API_KEY is set for vision transcription."
        )
    return text
